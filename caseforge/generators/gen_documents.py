#!/usr/bin/env python3
# =====================================================================
# gen_documents.py  —  Alltagsdokumente (Noise) auf allen drei Geraeten
# ---------------------------------------------------------------------
# Legt realistische Downloads/Dokumente (docx/xlsx/csv/txt/pdf) in die
# typischen Verzeichnisse je OS. Ueberwiegend NOISE; wenige Dateien sind
# dezent fallrelevant (Daniel: Kredit/Werkstattrechnung -> Schuldenmotiv;
# Anna: Anwalt-Infoblatt/Wohnungsexpose -> Trennungsabsicht).
# Schreibt ein Manifest mit Relevanz-Einordnung (fuer Dozenten).
# =====================================================================
import os
import csv
from datetime import datetime
from docx import Document
import openpyxl
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import cm
from reportlab.pdfgen import canvas

HERE = os.path.dirname(os.path.abspath(__file__))
BUILD = os.path.dirname(HERE)
ROOT = os.path.join(os.path.dirname(BUILD), "examples", "operation_waldweg")

import sys
sys.path.insert(0, HERE)
import case_master_io as cmio
import caseforge_rng as cfr
import noise_pools as npool

# Fall-Root aus Env ableiten (CaseForge), sonst Referenzfall
IOS = os.environ.get('WALDWEG_IOS_FS', os.path.join(ROOT, '01_ios_full_fs'))
AND = os.environ.get('WALDWEG_AND_FS', os.path.join(ROOT, '02_android_full_fs'))
WIN = os.environ.get('WALDWEG_WIN_FS', os.path.join(ROOT, '03_windows_triage'))
WUSER = cmio.windows_username()   # Windows-Profilordner aus Fall-Besitzer
CASE_ROOT = os.path.dirname(IOS) if os.environ.get('WALDWEG_IOS_FS') else ROOT

# Typische Zielverzeichnisse je Geraet
IOS_DL = os.path.join(IOS, 'private/var/mobile/Library/Mobile Documents/com~apple~CloudDocs/Downloads')
IOS_DOC = os.path.join(IOS, 'private/var/mobile/Library/Mobile Documents/com~apple~CloudDocs/Dokumente')
AND_DL = os.path.join(AND, 'storage/emulated/0/Download')
AND_DOC = os.path.join(AND, 'storage/emulated/0/Documents')
WIN_DL = os.path.join(WIN, f'C/Users/{WUSER}/Downloads')
WIN_DOC = os.path.join(WIN, f'C/Users/{WUSER}/Documents')

manifest = []  # (geraet, pfad_rel, relevanz, beschreibung)


def reg(geraet, path, relevanz, beschr):
    manifest.append((geraet, os.path.relpath(path, CASE_ROOT), relevanz, beschr))


def ensure(d):
    os.makedirs(d, exist_ok=True)


def w_txt(path, text):
    ensure(os.path.dirname(path))
    with open(path, 'w', encoding='utf-8') as f:
        f.write(text)


def w_csv(path, header, rows):
    ensure(os.path.dirname(path))
    with open(path, 'w', newline='', encoding='utf-8') as f:
        wr = csv.writer(f); wr.writerow(header); wr.writerows(rows)


def w_docx(path, title, paragraphs):
    ensure(os.path.dirname(path))
    doc = Document(); doc.add_heading(title, level=1)
    for p in paragraphs:
        doc.add_paragraph(p)
    doc.save(path)


def w_xlsx(path, sheet, header, rows):
    ensure(os.path.dirname(path))
    wb = openpyxl.Workbook(); ws = wb.active; ws.title = sheet
    ws.append(header)
    for r in rows:
        ws.append(r)
    wb.save(path)


def w_pdf(path, title, lines):
    ensure(os.path.dirname(path))
    c = canvas.Canvas(path, pagesize=A4); W, H = A4
    c.setFont("Helvetica-Bold", 15); c.drawString(2*cm, H-2.5*cm, title)
    c.setFont("Helvetica", 10); y = H-3.5*cm
    for ln in lines:
        c.drawString(2*cm, y, ln); y -= 0.6*cm
    c.showPage(); c.save()


# =====================================================================
def build_ios():
    # --- Noise ---
    w_pdf(os.path.join(IOS_DL, "Kursplan_Studio_Januar.pdf"), "Kursplan Januar 2026",
          ["Mo 18:00 Power Yoga", "Di 17:30 Ruecken Fit", "Do 19:00 Faszientraining",
           "Sa 10:00 Lauftreff", "Hinweis: Feiertage geschlossen."])
    reg("iPhone", os.path.join(IOS_DL, "Kursplan_Studio_Januar.pdf"), "noise", "Fitnessstudio-Kursplan (Download)")
    w_csv(os.path.join(IOS_DOC, "einkaufsliste.csv"), ["artikel", "menge"],
          [["Milch", "2"], ["Brot", "1"], ["Apfel", "6"], ["Kaffee", "1"], ["Spülmittel", "1"]])
    reg("iPhone", os.path.join(IOS_DOC, "einkaufsliste.csv"), "noise", "Einkaufsliste")
    w_pdf(os.path.join(IOS_DL, "Rezept_Ofengemuese.pdf"), "Rezept: Mediterranes Ofengemuese",
          ["Zutaten: Zucchini, Paprika, Aubergine, Olivenoel, Kraeuter",
           "Zubereitung: Bei 200 Grad ca. 30 Minuten backen."])
    reg("iPhone", os.path.join(IOS_DL, "Rezept_Ofengemuese.pdf"), "noise", "Rezept-PDF")
    w_txt(os.path.join(IOS_DOC, "packliste_ausflug.txt"),
          "Packliste Wochenende:\n- Wanderschuhe\n- Regenjacke\n- Snacks fuer Ben\n- Powerbank\n")
    reg("iPhone", os.path.join(IOS_DOC, "packliste_ausflug.txt"), "noise", "Packliste")
    # --- dezent relevant (Trennungsabsicht) ---
    w_pdf(os.path.join(IOS_DL, "Infoblatt_Familienrecht_Trennung.pdf"),
          "Infoblatt Familienrecht - Trennung & Sorgerecht",
          ["Erstberatung: Ablauf einer einvernehmlichen Trennung.",
           "Sorgerecht: gemeinsames vs. alleiniges Sorgerecht.",
           "Unterhalt: Kindes- und Trennungsunterhalt.",
           "Hinweis: ersetzt keine individuelle Rechtsberatung."])
    reg("iPhone", os.path.join(IOS_DL, "Infoblatt_Familienrecht_Trennung.pdf"), "context",
        "Anwalts-/Familienrecht-Infoblatt -> stuetzt Trennungsabsicht (passt zu Safari-Suche 07:05)")
    w_docx(os.path.join(IOS_DOC, "Wohnung_Expose_2Zimmer.docx"), "Expose - 2-Zimmer-Wohnung",
           ["Helle 2-Zimmer-Wohnung, 58 qm, Nachbarstadt.",
            "Kaltmiete 720 EUR, ab 01.03. verfuegbar.",
            "Notiz: Besichtigungstermin anfragen."])
    reg("iPhone", os.path.join(IOS_DOC, "Wohnung_Expose_2Zimmer.docx"), "context",
        "Wohnungsexpose -> konkrete Auszugsplanung (passt zu Suche 'wohnung mieten')")


def build_android():
    # --- Noise ---
    w_pdf(os.path.join(AND_DL, "Tankquittung_2026-01-20.pdf"), "Tankquittung",
          ["Station: Aral Stuttgart-Ost", "Datum: 20.01.2026 07:42",
           "Diesel 41,3 L  -  68,90 EUR", "Zahlung: EC-Karte"])
    reg("Samsung", os.path.join(AND_DL, "Tankquittung_2026-01-20.pdf"), "noise", "Tankquittung")
    w_xlsx(os.path.join(AND_DOC, "Aussendienst_Spesen_Jan2026.xlsx"), "Spesen",
           ["datum", "kunde", "ort", "km", "betrag_eur"],
           [["2026-01-14", "Klinik Sued", "Stuttgart", 42, 18.90],
            ["2026-01-20", "Praxis Dr. Weber", "Esslingen", 65, 29.30],
            ["2026-01-22", "Reha-Zentrum", "Goeppingen", 88, 39.60]])
    reg("Samsung", os.path.join(AND_DOC, "Aussendienst_Spesen_Jan2026.xlsx"), "noise", "Spesenabrechnung Aussendienst")
    w_txt(os.path.join(AND_DL, "navi_route_notiz.txt"),
          "Routenidee Wochenende: A8 Richtung Albtrauf, Wanderparkplatz.\nDauer ca. 50 min.\n")
    reg("Samsung", os.path.join(AND_DL, "navi_route_notiz.txt"), "noise", "Routennotiz")
    w_csv(os.path.join(AND_DOC, "sportapp_export.csv"), ["datum", "schritte", "km"],
          [["2026-01-22", "8421", "6.1"], ["2026-01-23", "5110", "3.7"], ["2026-01-24", "9302", "6.8"]])
    reg("Samsung", os.path.join(AND_DOC, "sportapp_export.csv"), "noise", "Schritt-/Sport-Export")
    # --- dezent relevant (Schuldenmotiv) ---
    w_pdf(os.path.join(AND_DL, "Rechnung_Werkstatt_Klenk_7711.pdf"),
          "Rechnung Nr. 7711 - Kfz-Werkstatt Klenk",
          ["Kunde: D. Reuter", "Leistung: Bremsen vorne/hinten, Inspektion",
           "Betrag: 1.480,00 EUR", "Zahlungsziel: 14 Tage - faellig 15.01.2026",
           "Hinweis: 2. Mahnung bei Verzug."])
    reg("Samsung", os.path.join(AND_DL, "Rechnung_Werkstatt_Klenk_7711.pdf"), "context",
        "Werkstattrechnung Klenk offen -> stuetzt Schulden/Glaeubiger-Motiv (passt zu WhatsApp-Ultimatum)")
    w_docx(os.path.join(AND_DOC, "Privatkredit_Vergleich.docx"), "Notizen Privatkredit",
           ["Anbieter A: 10.000 EUR, 4,9% eff., 48 Monate.",
            "Anbieter B: Sofortkredit, schnelle Auszahlung, hoehere Zinsen.",
            "Notiz: kurzfristig Liquiditaet noetig."])
    reg("Samsung", os.path.join(AND_DOC, "Privatkredit_Vergleich.docx"), "context",
        "Kreditvergleich-Notiz -> finanzieller Druck (passt zu Chrome/Edge-Suche)")


def build_windows():
    # --- Noise ---
    w_xlsx(os.path.join(WIN_DOC, "Haushaltsbudget_2026.xlsx"), "Budget",
           ["posten", "monatlich_eur"],
           [["Miete", 1180], ["Strom/Gas", 210], ["Lebensmittel", 620],
            ["Versicherungen", 340], ["Auto", 290], ["Sonstiges", 250]])
    reg("Notebook", os.path.join(WIN_DOC, "Haushaltsbudget_2026.xlsx"), "context",
        "Haushaltsbudget -> finanzielle Anspannung erkennbar (Noise/Context)")
    w_pdf(os.path.join(WIN_DL, "Bedienungsanleitung_Router.pdf"), "Kurzanleitung WLAN-Router",
          ["1. Geraet anschliessen", "2. WLAN-Namen waehlen", "3. Passwort vergeben",
           "Support: hotline@example-isp.de"])
    reg("Notebook", os.path.join(WIN_DL, "Bedienungsanleitung_Router.pdf"), "noise", "Router-Anleitung")
    w_txt(os.path.join(WIN_DOC, "todo.txt"),
          "ToDo:\n- Reifen wechseln lassen\n- Geschenk Ben Geburtstag\n- Steuerunterlagen sortieren\n- Klenk anrufen\n")
    reg("Notebook", os.path.join(WIN_DOC, "todo.txt"), "context",
        "ToDo-Liste -> beilaeufige Notiz 'Klenk anrufen' (Glaeubiger)")
    w_csv(os.path.join(WIN_DL, "kontoauszug_export.csv"), ["datum", "buchungstext", "betrag_eur"],
          [["2026-01-05", "Gehalt", "3120.00"], ["2026-01-08", "Miete", "-1180.00"],
           ["2026-01-12", "Supermarkt", "-86.40"], ["2026-01-15", "Ueberweisung T. Klenk", "-300.00"],
           ["2026-01-22", "Tankstelle", "-68.90"]])
    reg("Notebook", os.path.join(WIN_DL, "kontoauszug_export.csv"), "context",
        "Kontoauszug -> Teilzahlung an T. Klenk 300 EUR (Schuldenkontext)")
    w_docx(os.path.join(WIN_DOC, "Versicherungen_Uebersicht.docx"), "Versicherungen - Uebersicht",
           ["Hausrat: Police 4471", "Kfz-Haftpflicht: Police 8123",
            "Risikolebensversicherung: Police 9920 (Beguenstigte: Ehepartner)."])
    reg("Notebook", os.path.join(WIN_DOC, "Versicherungen_Uebersicht.docx"), "context",
        "Versicherungsuebersicht -> Risikolebensversicherung (passt zu Edge-Suche 'auszahlung todesfall')")


DEVBASE = {
    "ios":     {"downloads": IOS_DL, "documents": IOS_DOC},
    "android": {"downloads": AND_DL, "documents": AND_DOC},
    "windows": {"downloads": WIN_DL, "documents": WIN_DOC},
}
DEVLABEL = {"ios": "iPhone", "android": "Samsung", "windows": "Notebook"}


def build_from_master(docs):
    """Erzeugt Dokumente aus master.documents. Felder: device, area, name,
    kind(txt|csv|docx|xlsx|pdf), relevance, desc, lines|text|header+rows."""
    for d in docs:
        dev = (d.get("device") or "").lower()
        area = (d.get("area") or "downloads").lower()
        base = DEVBASE.get(dev, {}).get(area)
        if not base or not d.get("name"):
            continue
        path = os.path.join(base, d["name"])
        kind = (d.get("kind") or "txt").lower()
        title = d.get("title") or os.path.splitext(d["name"])[0]
        lines = d.get("lines") or ([d["text"]] if d.get("text") else [])
        if kind == "txt":
            w_txt(path, d.get("text") or "\n".join(lines) + "\n")
        elif kind == "csv":
            w_csv(path, d.get("header", []), d.get("rows", []))
        elif kind == "xlsx":
            w_xlsx(path, d.get("sheet", "Tabelle1"), d.get("header", []), d.get("rows", []))
        elif kind == "docx":
            w_docx(path, title, lines)
        elif kind == "pdf":
            w_pdf(path, title, lines)
        else:
            continue
        reg(DEVLABEL.get(dev, dev), path, d.get("relevance", "noise"), d.get("desc", ""))


def write_manifest():
    out = os.path.join(CASE_ROOT, '06_master', 'Dokumente_Manifest.csv')
    ensure(os.path.dirname(out))
    with open(out, 'w', newline='', encoding='utf-8') as f:
        wr = csv.writer(f)
        wr.writerow(["geraet", "pfad", "relevanz", "beschreibung"])
        wr.writerows(sorted(manifest))
    return out


def main():
    docs = cmio.documents()
    if docs:
        print(f"Dokumente-Inhaltsquelle: Master ({len(docs)} Dokumente)")
        build_from_master(docs)
    elif cfr.is_reference():
        print("Dokumente-Inhaltsquelle: Referenz-Fallback")
        build_ios(); build_android(); build_windows()
    else:
        # seed-gezogene, scope-skalierte Noise-Dokumente aus dem Pool
        lang = cmio.language_short()
        n = cmio.noise_count(10, key="documents")
        pool = npool.docs(lang)
        picked = cfr.sample(pool, n, salt="docnoise")
        devs = ["ios", "android", "windows"]
        gen = []
        for i, (name, kind, desc) in enumerate(picked):
            dev = devs[i % len(devs)]
            gen.append({"device": dev, "area": "downloads" if i % 2 else "documents",
                        "name": f"{i:02d}_{name}", "kind": kind,
                        "relevance": "noise", "desc": desc,
                        "lines": [desc], "text": desc,
                        "header": ["pos", "wert"], "rows": [["a", "1"], ["b", "2"]]})
        print(f"Dokumente-Inhaltsquelle: Pool/seed (scope, {len(gen)} Dokumente)")
        build_from_master(gen)
    out = write_manifest()
    n = len(manifest)
    rel = sum(1 for m in manifest if m[2] == "context")
    print(f"Erzeugt: {n} Dokumente ({n-rel} noise, {rel} context)")
    print(f"Manifest: {os.path.relpath(out, CASE_ROOT)}")
    for g, p, r, b in sorted(manifest):
        print(f"  [{r:7s}] {g:9s} {os.path.basename(p)}")


if __name__ == "__main__":
    main()
